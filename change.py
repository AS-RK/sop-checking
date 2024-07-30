import streamlit as st
from groq import Groq
import os
import base64
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from google.auth.transport.requests import Request
from datetime import datetime, timedelta
from email.mime.text import MIMEText
import pdfplumber
from io import BytesIO
from docx import Document
import pandas as pd
import json

# Set up the necessary scopes and credentials file
SCOPES = ['https://www.googleapis.com/auth/gmail.readonly', 'https://www.googleapis.com/auth/gmail.send']
GOOGLE_CREDENTIALS_JSON = st.secrets['GOOGLE_CREDENTIALS_JSON']

def get_gmail_service():
    if 'creds' not in st.session_state:
        st.session_state.creds = None

    creds = st.session_state.creds

    if creds:
        creds = Credentials.from_authorized_user_info(json.loads(creds), SCOPES)

    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            credentials_info = json.loads(GOOGLE_CREDENTIALS_JSON)
            flow = InstalledAppFlow.from_client_config(credentials_info, SCOPES, redirect_uri='https://xbtyidczfte6j8nkwmvhm4.streamlit.app/')

            auth_url, _ = flow.authorization_url(prompt='consent')

            st.write("Please go to this URL to authorize the application:")
            st.write(auth_url)

            auth_code = st.experimental_get_query_params().get('code')
            if auth_code:
                auth_code = auth_code[0]
                flow.fetch_token(code=auth_code)
                creds = flow.credentials
                st.session_state.creds = creds.to_json()

        if creds:
            st.session_state.creds = creds.to_json()

    service = build('gmail', 'v1', credentials=creds)
    return service

def get_message_details(message):
    """Extracts the subject and content of an email message."""
    headers = message['payload']['headers']
    subject = next(header['value'] for header in headers if header['name'] == 'Subject')
    
    if 'data' in message['payload']['body']:
        body_data = message['payload']['body']['data']
        body_decoded = base64.urlsafe_b64decode(body_data).decode('utf-8')
        return subject, body_decoded
    else:
        parts = message['payload'].get('parts', [])
        for part in parts:
            if part['mimeType'] == 'text/plain':
                body_data = part['body'].get('data', '')
                body_decoded = base64.urlsafe_b64decode(body_data).decode('utf-8')
                return subject, body_decoded
            elif part['mimeType'] == 'text/html':
                body_data = part['body'].get('data', '')
                body_decoded = base64.urlsafe_b64decode(body_data).decode('utf-8')
                return subject, body_decoded
    return subject, "No content available"

def fetch_gmail(sender_email):

    date = st.date_input('Date', key='date')
    message_content = ''
    subject = ''

    if st.button('Fetch Gmail Messages', key='fetch_gmail'):
        if sender_email and date:
            service = get_gmail_service()
            date_str = date.strftime('%Y/%m/%d')
            next_date_str = (date + timedelta(days=1)).strftime('%Y/%m/%d')
            query = f'from:{sender_email} after:{date_str} before:{next_date_str}'
            results = service.users().messages().list(userId='me', q=query).execute()
            messages = results.get('messages', [])

            if not messages:
                st.write('No messages found from this sender on the specified date.')
                st.session_state.gmail_fetched = False
            else:
                st.write(f'Found {len(messages)} messages from {sender_email} on {date}:')

                for msg in messages[::-1]:
                    msg_id = msg['id']
                    message = service.users().messages().get(userId='me', id=msg_id).execute()
                    snippet = message['snippet']
                    subject, message_content = get_message_details(message)
                    
                    st.write('---')
                    st.write(f'Subject: {subject}')
                    st.write(f'Message ID: {msg_id}')
                    st.write(f'Snippet: {snippet}')
                    st.write('Content:')
                    st.markdown(message_content, unsafe_allow_html=True)
                    st.write('---')
                
                st.session_state.gmail_fetched = True
                st.session_state.gmail_content = f"""Subject:{subject}\nContent:{message_content}"""
                
        else:
            st.write('Please enter a sender email address and a date.')
            st.session_state.gmail_fetched = False

def create_message(sender, to, subject, message_text):
    """Create a message for an email."""
    message = MIMEText(message_text)
    message['to'] = to
    message['from'] = sender
    message['subject'] = subject
    raw_message = base64.urlsafe_b64encode(message.as_bytes()).decode()
    return {'raw': raw_message}

def send_message(service, user_id, message):
    """Send an email message."""
    try:
        sent_message = service.users().messages().send(userId=user_id, body=message).execute()
        st.write(f"Message sent successfully: {sent_message['id']}")
    except Exception as error:
        st.write(f"An error occurred: {error}")

def gmailsender():
    st.title('Send an Email via Gmail')
    st.write('Enter the details below to send an email.')

    sender_email = st.text_input('Sender Email Address')
    recipient_email = st.text_input('Recipient Email Address')
    subject = st.text_input('Subject')
    message_text = st.text_area('Message',height=500)

    if st.button('Send Email'):
        if sender_email and recipient_email and subject and message_text:
            service = get_gmail_service()
            message = create_message(sender_email, recipient_email, subject, message_text)
            send_message(service, 'me', message)
        else:
            st.write('Please fill out all fields.')

def read_pdf(file):
    content = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            tables = page.extract_tables()
            if text:
                content.append(text)
            if tables:
                for table in tables:
                    table_text = "\n".join(["\t".join(map(str, row)) for row in table])
                    content.append(table_text)
    return "\n\n".join(content)

def read_docx(file):
    content = []
    document = Document(file)
    
    # Read paragraphs
    for paragraph in document.paragraphs:
        if paragraph.text.strip():  # Only add non-empty paragraphs
            content.append(paragraph.text)
    
    # Read tables
    for table in document.tables:
        for row in table.rows:
            row_text = "\t".join([cell.text.strip() for cell in row.cells])
            if row_text.strip():  # Only add non-empty rows
                content.append(row_text)
    
    # Join content with two newlines, but ensure no multiple consecutive newlines
    return "\n\n".join([line for line in content if line.strip()])

def parse_feedback(text):
    feedback_part = text.split("**Evaluation Based on SOP Criteria:**")[0]
    evaluation_part = text.split("**Evaluation Based on SOP Criteria:**")[1]
    
    criteria = []
    marks = []
    reasons = []
    
    for line in evaluation_part.split('\n'):
        if line.strip():
            crit, rest = line.split(":", 1)
            mark, reason = rest.split("(", 1)
            criteria.append(crit.strip())
            marks.append(mark.strip())
            reasons.append(reason.rstrip(")").strip())
    
    return feedback_part, criteria, marks, reasons

# Function to process the input text
def process_feedback(text):
    feedback, sop_evaluation = text.split("**Evaluation Based on SOP:**")
    return feedback.strip(), sop_evaluation.strip()

# Function to convert SOP evaluation to DataFrame
def parse_sop_evaluation(sop_text):
    lines = sop_text.split('\n')
    criteria = []
    marks = []
    reasons = []

    for line in lines:
        if '|' in line and 'Criteria' not in line:
            parts = line.split('|')
            criteria.append(parts[1].strip())
            marks.append(parts[2].strip())
            reasons.append(parts[3].strip())

    data = {
        'Criteria': criteria,
        'Mark (out of 10)': marks,
        'Reason': reasons
    }

    return pd.DataFrame(data)

def evaluator(client):
    st.title("Step 1: Upload SOP File in any one of the file format")
    uploaded_file = st.file_uploader("Choose a text file", type="txt")
    uploaded_file_pdf = st.file_uploader("Choose a PDF file", type="pdf")
    uploaded_file_docx = st.file_uploader("Choose a Word document", type="docx")
    if uploaded_file is not None:
        sop_content = uploaded_file.read().decode("utf-8")
        st.session_state.sop_uploaded = True
        modify_sop_content = st.text_area("SOP",sop_content, height=300)
        st.session_state.sop_content = modify_sop_content
    elif uploaded_file_docx is not None:
        sop_content = read_docx(uploaded_file_docx)
        st.session_state.sop_uploaded = True
        modify_sop_content = st.text_area("SOP",sop_content, height=300)
        st.session_state.sop_content = modify_sop_content
    elif uploaded_file_pdf is not None:
        sop_content = read_pdf(uploaded_file_pdf)
        st.session_state.sop_uploaded = True
        modify_sop_content = st.text_area("SOP",sop_content, height=300)
        st.session_state.sop_content = modify_sop_content
    else:
        sop_content = st.text_area("Type your SOP content or Modify existing sop content", height=300)
    
    if st.button("Insert text"):
        st.session_state.sop_uploaded = True
        st.session_state.sop_content = sop_content

    if st.session_state.sop_uploaded:
        st.title("Step 2: Client request")
        option = st.selectbox("Choose the way you want get client request", ("By typing","By gmail",), index=0, placeholder='Choose an option')
        if option == 'By gmail':
            st.write('Enter the sender email address and the date to fetch your Gmail messages from that sender.')
            sender_email = st.text_input('Sender Email Address', key='sender_email')
            fetch_gmail(sender_email)
        else:
            client_request = st.text_area("Client Request:", height=500)
            if st.button("Insert Request"):
                st.session_state.gmail_content = client_request
                st.session_state.gmail_fetched = True

        if st.session_state.get('gmail_fetched', False):
            st.title("Step 3: Type your content to evaluate")
            user_input = st.text_area("Your content:", height=400)
            if st.button("Evaluate"):
                if len(user_input) < 20:
                    st.error("Insufficient Information")
                else:
                    prompt = f"""
                                As a Quality Analyst, your task is to meticulously evaluate a user's response to a client email based on 
                                our Standard Operating Procedure (SOP) for email communication. The client email outlines an issue or concern 
                                they are experiencing with our product. Your evaluation involves identifying the specific problem mentioned by 
                                the client and ensuring the response adheres to our SOP. Follow these steps:
                                
                                SOP Content
                                {st.session_state.sop_content}
                                
                                Client Email
                                {st.session_state.gmail_content}
                                
                                Evaluation Task
                                Client's Issue:
                                
                                Clearly identify the specific problem or concern mentioned by the client in their email.
                                Constructive Feedback:
                                
                                Provide actionable feedback aimed at improving future responses.
                                Ensure feedback is specific and provides clear examples where applicable.
                                Evaluation Based on SOP:
                                
                                For each criterion in the SOP, provide a mark (out of 10) with a reason for the score within 25 words.
                                Present the criteria in a 2D list format: [[criteria], [mark(out of 10)], [reason]].
                                Suggested Alternatives:
                                
                                Suggest better alternative email content, fully structured with subject and body, 
                                that aligns with the SOP and addresses the client's concern effectively.
                        """

                    try:
                        completion = client.chat.completions.create(
                            messages=[
                                {"role": "system", "content": prompt},
                                {"role": "user", "content": user_input}
                            ],
                            model="llama3-8b-8192",
                            temperature=0,
                        )
                        st.session_state.feedback = completion.choices[0].message.content
                        
                        
                        # Call gmailsender() as Step 4
                        
                    except Exception as e:
                        st.error(f"An error occurred: {e}")
            if st.session_state.feedback:
                # Split the feedback into two parts: before and after the suggested alternatives
                feedback_parts = st.session_state.feedback.split("**Suggested Alternatives:**")
                feedback_text = feedback_parts[0].strip()
                # if feedback_text:
                #     feedback, criteria, marks, reasons = parse_feedback(feedback_text)
                    
                #     st.subheader("Feedback")
                #     st.write(feedback)
                    
                #     st.subheader("Evaluation Based on SOP Criteria")
                #     evaluation_data = pd.DataFrame({
                #         "Criteria": criteria,
                #         "Evaluation Mark": marks,
                #         "Reason": reasons
                #     })
                #     evaluation_data.index = evaluation_data.index + 1  # Adjust index to start from 1
                #     st.table(evaluation_data)

                if feedback_text:
                    feedback, sop_evaluation = process_feedback(feedback_text)
                    
                    st.subheader('Feedback')
                    st.write(feedback)
                    
                    st.subheader('Evaluation Based on SOP')
                    df = parse_sop_evaluation(sop_evaluation)
                    st.table(df)
                
                suggested_alternatives_text = feedback_parts[1].strip()

                # Further split the suggested alternatives into subject and content
                # Ensuring robust extraction by locating the "Subject:" and "Dear Jane," occurrences
                subject_start = suggested_alternatives_text.find("Subject:")
                subject_end = suggested_alternatives_text.find("\n\n", subject_start)
                subject = suggested_alternatives_text[subject_start + len("Subject:"):subject_end].strip()

                # Extracting the content part, ensuring that it starts right after the subject section
                content_start = subject_end + 2
                content = suggested_alternatives_text[content_start:].strip()

                # Streamlit app layout
                st.title("Suggested Alternatives")

                # Display feedback text area
                # st.subheader("Feedback")
                # st.text_area("feedback", feedback_text, height=500)
                # Display suggested alternatives
                st.subheader("Suggested Alternatives")
            
                st.text_area("Subject", subject, height=100)
                st.text_area("Content", content, height=300)
            if st.button("Step 4: Send Email") or st.session_state.gmail_send:
                st.session_state.gmail_send = True
                gmailsender()

def main():
    if 'sop_uploaded' not in st.session_state:
        st.session_state.sop_uploaded = False
    if 'gmail_fetched' not in st.session_state:
        st.session_state.gmail_fetched = False
    if 'gmail_content' not in st.session_state:
        st.session_state.gmail_content = ""
    if 'gmail_send' not in st.session_state:
        st.session_state.gmail_send = False
    if 'feedback' not in st.session_state:
        st.session_state.feedback = ""

    client = Groq(api_key=st.secrets["API_KEY"])
    option = st.selectbox("Choose the tool", ("Evaluator",), index=None, placeholder='Choose an option')
    if option == "Evaluator":
        evaluator(client)

if __name__ == "__main__":
    main()
